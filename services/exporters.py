from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from i18n import get_translations


def _flat_row(item: dict[str, Any]) -> dict[str, Any]:
    sources = "; ".join(
        f"{s.get('source', '')}: {s.get('source_url', '')}"
        for s in item.get("evidence", [])
    )
    resources = item.get("resource_estimate", {})
    return {
        "hypothesis_id": item["id"],
        "hypothesis": item["hypothesis"],
        "mechanism": item.get("mechanism", ""),
        "expected_effect": item.get("expected_effect", ""),
        "industrial_scale": item.get("industrial_scale", ""),
        "novelty_score": item["novelty_score"],
        "risk_score": item["risk_score"],
        "value_score": item["value_score"],
        "economic_value_score": item["economic_value_score"],
        "success_probability_score": item["success_probability_score"],
        "final_score": item["final_score"],
        "status": item["status"],
        "is_verified": item["is_verified"],
        "expert_rating": item.get("expert_rating"),
        "expert_comment": item.get("expert_comment", ""),
        "time_estimate": resources.get("time", ""),
        "cost_estimate": resources.get("cost", ""),
        "validation_volume": resources.get("volume", ""),
        "sources": sources,
    }


def _dicts_to_csv(rows: list[dict[str, Any]]) -> bytes:
    if not rows:
        return b""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=list(rows[0].keys()), delimiter=";")
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue().encode("utf-8-sig")


def build_csv(items: list[dict[str, Any]]) -> bytes:
    return _dicts_to_csv([_flat_row(item) for item in items])


def build_feedback_csv(rows: list[dict[str, Any]]) -> bytes:
    return _dicts_to_csv(rows)


def build_tasks_csv(run: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    rows = []
    for item in items:
        rows.append(
            {
                "Summary": f"Validate hypothesis {item['id']}",
                "Description": (
                    f"KPI: {run['kpi']}\n\n"
                    f"Hypothesis: {item['hypothesis']}\n\n"
                    f"Rationale: {item['rationale']}\n\n"
                    f"Mechanism: {item.get('mechanism', '')}\n\n"
                    f"Expected effect: {item.get('expected_effect', '')}\n\n"
                    f"Sources: " + "; ".join(s.get("source_url", "") for s in item.get("evidence", []))
                ),
                "Issue Type": "Task",
                "Priority": "High" if item["final_score"] >= 75 else "Medium",
                "Status": item["status"],
                "Hypothesis ID": item["id"],
            }
        )
    return _dicts_to_csv(rows)


def build_json(run: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    return json.dumps(
        {
            "kpi": run["kpi"],
            "constraints": run["constraints_text"],
            "language": run["language"],
            "knowledge_bases": run.get("knowledge_bases", []),
            "created_at": run["created_at"],
            "hypotheses": items,
        },
        ensure_ascii=False,
        indent=2,
    ).encode("utf-8")


def build_docx(run: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    tr = get_translations(run["language"])
    doc = Document()
    title = doc.add_heading(tr["report_title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{tr['kpi_label']}: {run['kpi']}")
    doc.add_paragraph(f"{tr['knowledge_bases_selected']}: {', '.join(tr.get(f'kb_{code}', code) for code in run.get('knowledge_bases', []))}")
    doc.add_paragraph(f"{tr['constraints']}: {run['constraints_text'] or tr['not_set']}")
    doc.add_paragraph(f"{tr['date']}: {run['created_at']}")

    doc.add_heading(tr["ranked_list"], level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["№", tr["rationale"], tr["novelty"], tr["risk"], tr["value"], tr["economic_value"], tr["final"]]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for rank, item in enumerate(items, start=1):
        values = [
            rank,
            item["hypothesis"],
            f"{item['novelty_score']:.2f}",
            f"{item['risk_score']:.2f}",
            f"{item['value_score']:.2f}",
            f"{item['economic_value_score']:.2f}",
            f"{item['final_score']:.2f}",
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)

    for rank, item in enumerate(items, start=1):
        doc.add_page_break()
        doc.add_heading(f"{rank}. {item['hypothesis']}", level=1)
        doc.add_paragraph(f"{tr['status']}: {item['status']}")
        doc.add_heading(tr["rationale"], level=2)
        doc.add_paragraph(item["rationale"])
        doc.add_heading(tr["mechanism"], level=2)
        doc.add_paragraph(item.get("mechanism", ""))
        doc.add_heading(tr["expected_effect"], level=2)
        doc.add_paragraph(item.get("expected_effect", ""))
        doc.add_heading(tr["industrial_scale"], level=2)
        doc.add_paragraph(item.get("industrial_scale", ""))
        doc.add_heading(tr["constraints_fit"], level=2)
        doc.add_paragraph(item.get("constraints_fit", ""))
        doc.add_heading(tr["verification"], level=2)
        doc.add_paragraph(item.get("verification_recommendation", ""))

        resources = item.get("resource_estimate", {})
        doc.add_heading(tr["resource_estimate"], level=2)
        doc.add_paragraph(
            f"{tr['time']}: {resources.get('time', '')}\n"
            f"{tr['cost']}: {resources.get('cost', '')}\n"
            f"{tr['volume']}: {resources.get('volume', '')}"
        )

        doc.add_heading(tr["sources"], level=2)
        for source in item.get("evidence", []):
            doc.add_paragraph(
                f"{tr['source']}: {source.get('source', '')}\n"
                f"{tr['source_link']}: {source.get('source_url', '')}\n"
                f"{tr['page']}: {source.get('page', '')}\n"
                f"{tr['chunk']}: {source.get('chunk_id', '')}\n"
                f"{source.get('evidence_fragment', '')}"
            )
        doc.add_heading(tr["expert_review"], level=2)
        doc.add_paragraph(
            f"{tr['hypothesis_verified']}: {bool(item.get('is_verified'))}\n"
            f"{tr['expert_rating']}: {item.get('expert_rating') or ''}\n"
            f"{tr['expert_comment']}: {item.get('expert_comment', '')}"
        )

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _font_for_language(language: str) -> tuple[str, bool]:
    if language == "zh":
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light", True
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/Library/Fonts/Arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("AppFont", str(path)))
            return "AppFont", False
    raise RuntimeError("A Unicode font was not found for PDF export.")


def build_pdf(run: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    tr = get_translations(run["language"])
    font_name, _ = _font_for_language(run["language"])
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontName=font_name, alignment=TA_CENTER, fontSize=17, leading=21)
    heading_style = ParagraphStyle("HeadingX", parent=styles["Heading1"], fontName=font_name, fontSize=13, leading=16)
    body_style = ParagraphStyle("BodyX", parent=styles["BodyText"], fontName=font_name, fontSize=9, leading=12)

    story = [
        Paragraph(escape(tr["report_title"]), title_style),
        Spacer(1, 8),
        Paragraph(escape(f"{tr['kpi_label']}: {run['kpi']}"), body_style),
        Paragraph(escape(f"{tr['knowledge_bases_selected']}: {', '.join(tr.get(f'kb_{code}', code) for code in run.get('knowledge_bases', []))}"), body_style),
        Paragraph(escape(f"{tr['constraints']}: {run['constraints_text'] or tr['not_set']}"), body_style),
        Spacer(1, 10),
        Paragraph(escape(tr["ranked_list"]), heading_style),
    ]
    table_data = [["№", tr["rationale"], "N", "R", "V", "EV", tr["final"]]]
    for rank, item in enumerate(items, start=1):
        table_data.append([
            str(rank),
            Paragraph(escape(item["hypothesis"]), body_style),
            f"{item['novelty_score']:.2f}",
            f"{item['risk_score']:.2f}",
            f"{item['value_score']:.2f}",
            f"{item['economic_value_score']:.2f}",
            f"{item['final_score']:.2f}",
        ])
    table = Table(table_data, colWidths=[8*mm, 93*mm, 13*mm, 13*mm, 13*mm, 15*mm, 18*mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), font_name),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E9EEF5")),
        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#98A2B3")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
    ]))
    story.extend([table, Spacer(1, 10)])

    for rank, item in enumerate(items, start=1):
        story.append(PageBreak())
        story.append(Paragraph(escape(f"{rank}. {item['hypothesis']}"), heading_style))
        sections = [
            (tr["rationale"], item.get("rationale", "")),
            (tr["mechanism"], item.get("mechanism", "")),
            (tr["expected_effect"], item.get("expected_effect", "")),
            (tr["industrial_scale"], item.get("industrial_scale", "")),
            (tr["constraints_fit"], item.get("constraints_fit", "")),
            (tr["verification"], item.get("verification_recommendation", "")),
        ]
        for heading, text in sections:
            story.append(Paragraph(f"<b>{escape(heading)}:</b> {escape(str(text))}", body_style))
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<b>{escape(tr['sources'])}</b>", body_style))
        for source in item.get("evidence", []):
            text = (
                f"{escape(source.get('source',''))}<br/>"
                f"{escape(tr['source_link'])}: {escape(source.get('source_url',''))}<br/>"
                f"{escape(source.get('evidence_fragment',''))}"
            )
            story.append(Paragraph(text, body_style))

    document.build(story)
    return output.getvalue()
